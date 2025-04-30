"""
This script creates two files:

1. Fishing_Calendar.docx – A formatted Word document with the updated 2025–2026 fishing calendar,
   including headings, tables, and a separate outline of fishing factors.
2. Fishing_Calendar.xlsx  – An Excel workbook with worksheets for each river location, each containing
   table data for the fishing calendar.

Required packages:
    - python-docx
    - openpyxl

Install them with:
    pip install python-docx openpyxl
"""

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

# -------------------------
# Data for the Fishing Calendar
# -------------------------

# For brevity, the data below is based on our updated calendar.
# Each table is represented as a list of rows (first row = header).

# Columbia River (Rainier/Prescott Beach) – Outgoing Tides Only
columbia_2025 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["02/15/2025", "10:30 AM", "Winter Steelhead – Peak", "45°F", "Falling", "First Quarter", "Outgoing @ 10:00 AM", "Pro-Cure Bloody Tuna, Anise Oil"],
    ["03/15/2025", "9:00 AM", "Spring Chinook – A Run (Starting)", "45–48°F", "Stable", "Full Moon", "Outgoing @ 9:30 AM", "Graybill's Tuna Belly, Cured Coon Shrimp"],
    ["04/05/2025", "11:00 AM", "Spring Chinook – Peak", "56°F", "Rising", "Waxing Crescent", "Outgoing @ 11:00 AM", "Pro-Cure Slam-Ola w/ Garlic"],
    ["04/12/2025", "12:30 PM", "Spring Steelhead – A Run (Starting)", "57°F", "Stable High", "First Quarter", "Outgoing @ 1:00 PM", "*Pro-Cure Slam-Ola Powder + Graybill's Tuna Belly + Cured Coon Shrimp*"],
    ["05/10/2025", "8:00 AM", "Spring Steelhead – B Run (Starting)", "54°F", "Falling", "Waning Crescent", "Outgoing @ 8:30 AM", "Sand Shrimp, Herring Oil"],
    ["06/20/2025", "12:00 PM", "Summer Chinook – B Run (Peak)", "64°F", "Stable", "Full Moon", "Outgoing @ 12:00 PM", "Pro-Cure Garlic, Krill Powder"],
    ["07/15/2025", "8:00 AM", "Summer Coho – A Run (Starting)", "66°F", "Rising", "Waning Gibbous", "Outgoing @ 9:00 AM", "Herring, Anchovy Oil"],
    ["08/05/2025", "7:30 AM", "Fall Chinook – A Run (Starting)", "68°F", "Falling", "First Quarter", "Outgoing @ 8:00 AM", "Tuna Oil, Sand Shrimp"],
    ["09/10/2025", "7:00 AM", "Fall Chinook – Peak", "66°F", "Stable", "Full Moon", "Outgoing @ 7:30 AM", "Pro-Cure Bloody Tuna, Sardine"],
    ["10/15/2025", "8:30 AM", "Coho – Peak", "59°F", "Rising", "Last Quarter", "Outgoing @ 9:00 AM", "Shrimp, Anise"],
    ["11/28/2025", "1:30 PM", "Winter Steelhead – B Run (Ending)", "48°F", "Falling", "New Moon", "Outgoing @ 1:00 PM", "Sardine, Herring Oil"],
    ["12/05/2025", "2:00 PM", "Winter Steelhead – A Run (Starting)", "46°F", "Stable", "Waxing Crescent", "Outgoing @ 2:30 PM", "Pro-Cure Garlic, Egg Cure"],
]
columbia_2026 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["01/25/2026", "12:45 PM", "Winter Steelhead – Peak", "42°F", "Stable", "First Quarter", "Outgoing @ 12:30 PM", "Pro-Cure Anise, Herring"],
    ["03/05/2026", "10:30 AM", "Spring Chinook – Early Season", "46°F", "Rising", "Waxing Crescent", "Outgoing @ 11:00 AM", "Graybill's Tuna Belly, Sand Shrimp"],
    ["04/04/2026", "11:30 AM", "Spring Chinook – A Run (Starting)", "55°F", "Falling", "Waxing Crescent", "Outgoing @ 11:00 AM", "Pro-Cure Slam-Ola, Cured Coon Shrimp"],
    ["04/25/2026", "7:45 AM", "Spring Steelhead – Starting", "56°F", "Stable", "Last Quarter", "Outgoing @ 8:15 AM", "Pro-Cure Bloody Tuna, Krill Powder"],
]

# Nestucca River (Pacific City) – Both Incoming and Outgoing
nestucca_2025 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["01/12/2025", "11:00 AM", "Winter Steelhead – A Run (Starting)", "50°F", "Falling", "New Moon", "Incoming @ 11:00 AM", "Pro-Cure Bloody Tuna, Egg Cure"],
    ["02/10/2025", "1:00 PM", "Winter Steelhead – Peak", "49°F", "Stable", "Waxing Crescent", "Outgoing @ 1:30 PM", "Pro-Cure Anise, Garlic"],
    ["04/20/2025", "9:30 AM", "Spring Chinook – A Run (Starting)", "54°F", "Rising", "Waxing Crescent", "Outgoing @ 10:00 AM", "Pro-Cure Slam-Ola, Graybill's Tuna Belly"],
    ["05/25/2025", "6:30 AM", "Summer Steelhead – Early Season", "52°F", "Stable High", "Waning Crescent", "Incoming @ 7:00 AM", "Pro-Cure Garlic, Sand Shrimp"],
    ["08/20/2025", "6:30 AM", "Fall Chinook – A Run (Starting)", "62°F", "Falling", "Waning Crescent", "Outgoing @ 7:00 AM", "Herring Oil, Sardine"],
    ["09/15/2025", "7:30 AM", "Summer Coho – Peak", "56°F", "Rising", "Full Moon", "Incoming @ 7:00 AM", "Pro-Cure Bloody Tuna, Anchovy"],
    ["10/07/2025", "8:00 AM", "Fall Coho – B Run (Peak)", "55°F", "Stable", "Waning Gibbous", "Outgoing @ 8:30 AM", "Tuna Oil, Anise"],
    ["12/25/2025", "10:30 AM", "Winter Steelhead – A Run (Starting)", "44°F", "Falling", "Waxing Crescent", "Outgoing @ 11:00 AM", "Egg Cure, Pro-Cure Anise"],
]
nestucca_2026 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["01/13/2026", "11:05 AM", "Winter Steelhead – A Run (Starting)", "50°F", "Stable", "New Moon", "Incoming @ 11:00 AM", "Pro-Cure Bloody Tuna, Egg Cure"],
    ["02/11/2026", "1:10 PM", "Winter Steelhead – Peak", "49°F", "Rising", "Waxing Crescent", "Outgoing @ 1:20 PM", "Pro-Cure Anise, Garlic"],
    ["03/25/2026", "8:30 AM", "Winter Steelhead – Late Season", "47°F", "Falling", "Last Quarter", "Outgoing @ 9:00 AM", "Sand Shrimp, Egg Cure"],
    ["04/20/2026", "8:30 AM", "Spring Chinook – A Run (Starting)", "52°F", "Stable High", "Waning Gibbous", "Outgoing @ 9:00 AM", "Pro-Cure Slam-Ola, Graybill's Tuna Belly"],
]

# Cowlitz River – Both Incoming and Outgoing
cowlitz_2025 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["03/25/2025", "8:30 AM", "Winter Steelhead – Peak", "44°F", "Rising", "Last Quarter", "Outgoing @ 9:00 AM", "Pro-Cure Anise, Shrimp"],
    ["03/30/2025", "10:00 AM", "Spring Chinook – A Run (Starting)", "55°F", "Stable", "Last Quarter", "Incoming @ 10:00 AM", "Graybill's Tuna Belly, Pro-Cure Slam-Ola"],
    ["06/25/2025", "6:00 AM", "Summer Steelhead – A Run (Starting)", "58°F", "Falling", "Waning Crescent", "Outgoing @ 6:30 AM", "Pro-Cure Garlic, Krill"],
    ["07/05/2025", "11:00 AM", "Summer Chinook – Peak", "62°F", "Stable High", "Full Moon", "Outgoing @ 11:30 AM", "Sand Shrimp, Tuna Oil"],
    ["08/12/2025", "9:45 AM", "Summer Coho – B Run (Ending)", "64°F", "Rising", "Waning Gibbous", "Incoming @ 10:15 AM", "Herring Oil, Anise"],
    ["10/25/2025", "9:15 AM", "Coho – Peak", "54°F", "Falling", "New Moon", "Outgoing @ 9:45 AM", "Pro-Cure Bloody Tuna, Sardine"],
    ["12/15/2025", "2:30 PM", "Winter Steelhead – A Run (Starting)", "47°F", "Stable", "New Moon", "Outgoing @ 2:45 PM", "Egg Cure, Pro-Cure Anise"],
]
cowlitz_2026 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["02/05/2026", "1:30 PM", "Winter Steelhead – Building", "41°F", "Rising", "Full Moon", "Outgoing @ 2:00 PM", "Egg Cure, Pro-Cure Anise"],
    ["03/15/2026", "9:45 AM", "Winter Steelhead – Peak", "45°F", "Stable", "Full Moon", "Incoming @ 10:15 AM", "Pro-Cure Anise, Sand Shrimp"],
    ["03/31/2026", "10:00 AM", "Spring Chinook – A Run (Starting)", "55°F", "Falling", "Last Quarter", "Incoming @ 10:00 AM", "Graybill's Tuna Belly, Pro-Cure Slam-Ola"],
    ["04/15/2026", "9:15 AM", "Spring Chinook – Building", "53°F", "Rising", "Full Moon", "Incoming @ 9:45 AM", "Graybill's Tuna Belly, Cured Coon Shrimp"],
]

# Lewis River – Both Incoming and Outgoing
lewis_2025 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["01/20/2025", "1:15 PM", "Winter Chinook – A Run (Peak)", "48°F", "Stable", "New Moon", "Outgoing @ 1:30 PM", "Pro-Cure Bloody Tuna, Herring"],
    ["04/10/2025", "10:30 AM", "Spring Chinook – B Run (Peak)", "55°F", "Rising", "Waxing Crescent", "Incoming @ 10:30 AM", "Pro-Cure Slam-Ola, Graybill's Tuna Belly"],
    ["05/20/2025", "7:30 AM", "Spring Chinook – Late Season", "53°F", "Falling", "Last Quarter", "Outgoing @ 8:00 AM", "Sand Shrimp, Anise"],
    ["07/22/2025", "9:00 AM", "Summer Steelhead – B Run (Peak)", "63°F", "Stable High", "Full Moon", "Incoming @ 9:00 AM", "Pro-Cure Garlic, Krill"],
    ["09/20/2025", "6:45 AM", "Fall Chinook – Peak", "58°F", "Rising", "Waning Gibbous", "Outgoing @ 7:15 AM", "Tuna Oil, Sardine"],
    ["11/10/2025", "11:30 AM", "Coho – Late Season", "50°F", "Stable", "First Quarter", "Incoming @ 12:00 PM", "Pro-Cure Bloody Tuna, Anchovy"],
    ["12/20/2025", "11:00 AM", "Winter Steelhead – C Run (Ending)", "46°F", "Falling", "Waning Crescent", "Incoming @ 11:15 AM", "Egg Cure, Pro-Cure Anise"],
]
lewis_2026 = [
    ["Date", "Time", "Event Description", "Water Temp", "Barometric", "Moon Phase", "Tide Details", "Recommended Scents"],
    ["01/21/2026", "1:20 PM", "Winter Chinook – A Run (Peak)", "48°F", "Stable", "New Moon", "Outgoing @ 1:30 PM", "Pro-Cure Bloody Tuna, Herring"],
    ["02/20/2026", "11:45 AM", "Winter Steelhead – Peak", "45°F", "Rising", "New Moon", "Incoming @ 12:15 PM", "Egg Cure, Pro-Cure Anise"],
    ["04/10/2026", "10:45 AM", "Spring Chinook – Peak", "54°F", "Stable High", "First Quarter", "Outgoing @ 11:15 AM", "Pro-Cure Slam-Ola, Graybill's Tuna Belly"],
]

# -------------------------
# Word Document Creation (using python-docx)
# -------------------------

def create_word_document():
    document = Document()
    document.add_heading("2025–2026 COMPREHENSIVE SALMON & STEELHEAD FISHING CALENDAR", 0)
    document.add_paragraph("Organized by River Location with Enhanced Fishing Details")
    document.add_paragraph("--------------------------------------------------------------------------------")
    
    # Helper function to add a table
    def add_table(section_title, data):
        document.add_heading(section_title, level=1)
        table = document.add_table(rows=1, cols=len(data[0]))
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(data[0]):
            hdr_cells[i].text = header
        for row_data in data[1:]:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                row_cells[i].text = cell_text
        document.add_paragraph()  # add space after table

    # Columbia River Section (Combine 2025 & 2026)
    document.add_heading("Columbia River (Rainier/Prescott Beach) [Outgoing Tides Only]", level=1)
    document.add_heading("2025", level=2)
    add_table("Columbia 2025", columbia_2025)
    document.add_heading("2026", level=2)
    add_table("Columbia 2026", columbia_2026)
    
    # Nestucca River Section
    document.add_heading("Nestucca River (Pacific City, Oregon)", level=1)
    document.add_heading("2025", level=2)
    add_table("Nestucca 2025", nestucca_2025)
    document.add_heading("2026", level=2)
    add_table("Nestucca 2026", nestucca_2026)
    
    # Cowlitz River Section
    document.add_heading("Cowlitz River", level=1)
    document.add_heading("2025", level=2)
    add_table("Cowlitz 2025", cowlitz_2025)
    document.add_heading("2026", level=2)
    add_table("Cowlitz 2026", cowlitz_2026)
    
    # Lewis River Section
    document.add_heading("Lewis River", level=1)
    document.add_heading("2025", level=2)
    add_table("Lewis 2025", lewis_2025)
    document.add_heading("2026", level=2)
    add_table("Lewis 2026", lewis_2026)
    
    # Add a separate section on Fishing Factors & Run Timing Details
    document.add_heading("Fishing Factors & Run Timing Details", level=1)
    fishing_factors = (
        "TIDE INFLUENCES:\n"
        "• Columbia River: Outgoing tides only (first 2 hours are optimal as fish concentrate).\n"
        "• Nestucca, Cowlitz, and Lewis Rivers: Both incoming (bring in fresh, oxygenated water) and outgoing tides (concentrate fish) are productive.\n\n"
        "WATER TEMPERATURE IMPACT:\n"
        "• 38–42°F: Fish lethargic; use slow presentations in deep holes.\n"
        "• 42–48°F: Moderate activity; slower presentations recommended.\n"
        "• 48–55°F: Ideal feeding temperatures for spring Chinook and steelhead.\n"
        "• 55–65°F: Peak activity; fast presentations are effective.\n"
        "• 65°F+: Fish may be stressed; early/late tactics advised.\n\n"
        "BAROMETRIC PRESSURE EFFECTS:\n"
        "• Falling pressure: Indicates an approaching storm; fish feed aggressively.\n"
        "• Rising pressure: Following storm fronts, feeding often resumes actively.\n"
        "• Stable high pressure: Produces predictable feeding behavior.\n\n"
        "MOON PHASES:\n"
        "• New Moon: Often increases fish movement in darkness.\n"
        "• Full Moon: May boost nocturnal feeding, with variable daytime activity.\n"
        "• First/Last Quarters: Transitional phases with variable effects.\n\n"
        "SCENT RECOMMENDATIONS:\n"
        "• Spring Chinook: Pro-Cure Slam-Ola Powder (with garlic), Graybill's Tuna Belly, and Cured Coon Shrimp.\n"
        "• Winter Steelhead: Pro-Cure Bloody Tuna and Pro-Cure Anise (or Egg Cure).\n"
        "• Summer/Fall: Sand Shrimp, Herring Oil, Tuna Oil, and Anchovy-based scents.\n\n"
        "RUN TIMING (GENERAL TRENDS):\n"
        "• Columbia River: Winter Steelhead (Dec–Mar), Spring Chinook (Mar–May), Spring Steelhead (Apr–Jun), Summer Chinook (Jun–Jul), Fall Coho/Chinook (Aug–Nov).\n"
        "• Nestucca River: Winter Steelhead (Dec–Apr), Spring Chinook (Apr–Jun), Summer Steelhead & Fall runs (May–Dec).\n"
        "• Cowlitz River: Winter Steelhead (Nov–Apr), Spring Chinook (Mar–Jun), Summer Steelhead/Chinook (Jun–Oct).\n"
        "• Lewis River: Winter Chinook/Steelhead (Dec–Mar), Spring Chinook (Apr–Jun), Summer Steelhead (Jun–Aug), Fall Chinook (late season).\n"
    )
    document.add_paragraph(fishing_factors)
    
    # Save the Word document.
    document.save("Fishing_Calendar.docx")
    print("Word document 'Fishing_Calendar.docx' created successfully.")

# -------------------------
# Excel File Creation (using openpyxl)
# -------------------------

def create_excel_workbook():
    wb = Workbook()
    
    # A helper function to add a worksheet given a title and table data.
    def add_worksheet(sheet_name, data):
        ws = wb.create_sheet(title=sheet_name)
        for row_idx, row in enumerate(data, start=1):
            for col_idx, cell_value in enumerate(row, start=1):
                ws.cell(row=row_idx, column=col_idx, value=cell_value)
        # Optional: set the column widths for readability.
        for i in range(1, len(data[0]) + 1):
            ws.column_dimensions[get_column_letter(i)].width = 20

    # Remove the default sheet created by openpyxl (if unused).
    if "Sheet" in wb.sheetnames:
        std = wb["Sheet"]
        wb.remove(std)
    
    # Add worksheets for each river and year (naming format: Location_Year).
    add_worksheet("Columbia_2025", columbia_2025)
    add_worksheet("Columbia_2026", columbia_2026)
    add_worksheet("Nestucca_2025", nestucca_2025)
    add_worksheet("Nestucca_2026", nestucca_2026)
    add_worksheet("Cowlitz_2025", cowlitz_2025)
    add_worksheet("Cowlitz_2026", cowlitz_2026)
    add_worksheet("Lewis_2025", lewis_2025)
    add_worksheet("Lewis_2026", lewis_2026)
    
    # Save the Excel workbook.
    wb.save("Fishing_Calendar.xlsx")
    print("Excel workbook 'Fishing_Calendar.xlsx' created successfully.")

# -------------------------
# Main Function
# -------------------------

def main():
    create_word_document()
    create_excel_workbook()

if __name__ == "__main__":
    main()